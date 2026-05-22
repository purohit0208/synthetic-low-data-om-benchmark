import json
import math
import shutil
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
SUB = ROOT / "submission" / "mdpi_ai_2026_05_22"
FIG = SUB / "figures_high_res"
SUPP = SUB / "supplementary"

ALPHA = chr(945)
DELTA = chr(916)
GAMMA = chr(947)
BETA = chr(946)
ETA = chr(951)
THETA = chr(952)
SIGMA = chr(963)
SUP2 = chr(178)
DOT = chr(183)
GE = chr(8805)
LE = chr(8804)
IN = chr(8712)
SUM = chr(931)
WEDGE = chr(8743)

COLORS = {
    "ink": "#1F2937",
    "muted": "#6B7280",
    "grid": "#E5E7EB",
    "blue": "#2563EB",
    "teal": "#0F766E",
    "green": "#16A34A",
    "amber": "#D97706",
    "red": "#DC2626",
    "purple": "#7C3AED",
    "slate": "#475569",
}

MODEL_LABELS = {
    "lr_struct": "Logistic\n(struct.)",
    "rf_struct": "Random Forest\n(struct.)",
    "xgb_struct": "XGBoost\n(struct.)",
    "lr_tfidf": "TF-IDF + LR\n(text)",
    "lr_minilm": "MiniLM + LR\n(text)",
    "fusion_tfidf": "Struct. +\nTF-IDF",
    "fusion_minilm": "Struct. +\nMiniLM",
}


def read_csv(name):
    return pd.read_csv(OUT / name)


def style_axes(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CBD5E1")
    ax.spines["bottom"].set_color("#CBD5E1")
    ax.tick_params(colors=COLORS["ink"], labelsize=9)
    ax.grid(True, color=COLORS["grid"], linewidth=0.7, axis="y")
    ax.set_axisbelow(True)


def save_fig(fig, name):
    path = FIG / name
    fig.savefig(path, dpi=600, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def fmt(x, digits=3):
    if isinstance(x, str):
        return x
    if pd.isna(x):
        return ""
    return f"{float(x):.{digits}f}"


def pct(x, digits=1):
    return f"{100 * float(x):.{digits}f}%"


def add_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def make_table(doc, headers, rows, widths=None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        hdr[i].text = h
        add_cell_shading(hdr[i], "EAF2F8")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Arial"
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Arial"
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.style = "Caption"
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    return p


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(10)
    return p


def set_doc_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Title", 18, "1F2937"),
        ("Heading 1", 14, "1F2937"),
        ("Heading 2", 12, "1F2937"),
        ("Heading 3", 10.5, "334155"),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    cap = styles["Caption"]
    cap.font.name = "Arial"
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(55, 65, 81)


def generate_figures():
    FIG.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({
        "font.family": "Arial",
        "axes.titlesize": 12,
        "axes.labelsize": 10,
        "legend.fontsize": 8.5,
        "figure.dpi": 150,
    })

    # Figure 1: benchmark pipeline.
    fig, ax = plt.subplots(figsize=(9.5, 4.6))
    ax.axis("off")
    boxes = [
        ("Latent degradation\nhidden simulator state", 0.08, 0.62, COLORS["blue"]),
        ("Noisy sensors\n+ sparse reports", 0.31, 0.62, COLORS["teal"]),
        ("Leakage audit\n+ fixed splits", 0.54, 0.62, COLORS["slate"]),
        ("Model stack\nstructured / text / fusion", 0.77, 0.62, COLORS["purple"]),
        ("Validation calibration\n+ threshold search", 0.54, 0.22, COLORS["amber"]),
        ("Test evaluation\nperformance, robustness,\nrouting, explanations", 0.77, 0.22, COLORS["green"]),
    ]
    for label, x, y, color in boxes:
        ax.text(
            x, y, label, ha="center", va="center", color=COLORS["ink"], fontsize=10,
            bbox=dict(boxstyle="round,pad=0.55,rounding_size=0.08", fc="white", ec=color, lw=2)
        )
    arrows = [((0.18, 0.62), (0.25, 0.62)), ((0.41, 0.62), (0.48, 0.62)),
              ((0.64, 0.62), (0.71, 0.62)), ((0.77, 0.52), (0.77, 0.33)),
              ((0.70, 0.22), (0.61, 0.22)), ((0.60, 0.31), (0.71, 0.53))]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.6, color=COLORS["muted"]))
    ax.set_title("Synthetic industrial O&M benchmark and evaluation flow", color=COLORS["ink"], pad=14, fontweight="bold")
    save_fig(fig, "figure_1_pipeline.png")

    # Figure 2: latent degradation and observed vibration.
    latent_path = OUT / "latent_database.parquet"
    if latent_path.exists():
        latent = pd.read_parquet(latent_path)
        data = pd.read_parquet(OUT / "dataset.parquet")
        selected = latent.groupby("asset_id")["latent_maintenance_triggered"].sum().sort_values(ascending=False).index[0]
        la = latent[latent["asset_id"] == selected].sort_values("date")
        da = data[data["asset_id"] == selected].sort_values("date")
        fig, axes = plt.subplots(2, 1, figsize=(8.2, 5), sharex=True)
        axes[0].plot(pd.to_datetime(la["date"]), la["latent_degradation"], color=COLORS["blue"], lw=1.8)
        axes[0].axhline(0.75, color=COLORS["red"], ls="--", lw=1.2, label="fault trigger")
        axes[0].set_ylabel("Latent D_t")
        axes[0].legend(loc="upper left", frameon=False)
        style_axes(axes[0])
        axes[1].plot(pd.to_datetime(da["date"]), da["vibration_rms"], color=COLORS["amber"], lw=1.6)
        axes[1].set_ylabel("Vibration RMS")
        axes[1].set_xlabel("Simulated date")
        style_axes(axes[1])
        fig.suptitle(f"Hidden degradation and noisy observable sensor response ({selected})", y=0.98, fontweight="bold")
        save_fig(fig, "figure_2_degradation.png")

    # Figure 3: learning curves.
    exp1 = read_csv("exp1_learning_curves.csv")
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    color_map = {
        "lr_struct": COLORS["slate"], "rf_struct": COLORS["green"], "xgb_struct": COLORS["red"],
        "lr_tfidf": COLORS["amber"], "lr_minilm": COLORS["purple"],
        "fusion_tfidf": COLORS["blue"], "fusion_minilm": COLORS["teal"],
    }
    for model, group in exp1.groupby("Model"):
        group = group.sort_values("Fraction")
        ax.plot(group["Fraction"] * 100, group["PR-AUC"], marker="o", lw=1.8, ms=4, label=model, color=color_map.get(model))
    ax.set_xlabel("Labeled training data (%)")
    ax.set_ylabel("Chronological-test PR-AUC")
    ax.set_title("Label-scarcity learning curves", fontweight="bold")
    style_axes(ax)
    ax.legend(ncol=2, frameon=False, loc="upper left")
    save_fig(fig, "figure_3_learning_curves.png")

    # Figure 4: random split vs held-out site.
    site = read_csv("table_7_cross_site_generalization.csv")
    models = list(MODEL_LABELS)
    fig, ax = plt.subplots(figsize=(8.8, 4.8))
    x = np.arange(len(models))
    ax.bar(x - 0.18, [site.loc[site.Model == m, "Random-Split Test PR-AUC"].iloc[0] for m in models],
           width=0.36, label="Random split", color=COLORS["blue"])
    ax.bar(x + 0.18, [site.loc[site.Model == m, "Cross-Site Mean PR-AUC"].iloc[0] for m in models],
           width=0.36, label="Held-out site mean", color=COLORS["red"])
    ax.set_xticks(x)
    ax.set_xticklabels([MODEL_LABELS[m] for m in models], rotation=0)
    ax.set_ylabel("PR-AUC")
    ax.set_title("Random splitting overstates cross-site generalization", fontweight="bold")
    style_axes(ax)
    ax.legend(frameon=False)
    save_fig(fig, "figure_4_site_shift.png")

    # Figure 5: robustness.
    robust = read_csv("exp3_robustness_results.csv")
    fig, axes = plt.subplots(2, 2, figsize=(9.2, 6.4))
    panels = [("Sensor Noise", axes[0, 0]), ("Missing Sensors", axes[0, 1]), ("Missing Reports", axes[1, 0]), ("Report Ambiguity", axes[1, 1])]
    show_models = ["rf_struct", "fusion_minilm", "lr_tfidf"]
    for perturb, ax in panels:
        subset = robust[robust["Perturbation_Type"] == perturb]
        for model in show_models:
            g = subset[subset["Model"] == model]
            ax.plot(range(len(g)), g["PR-AUC"], marker="o", lw=1.8, label=model, color=color_map[model])
        ax.set_xticks(range(len(g)))
        ax.set_xticklabels(list(g["Level"]))
        ax.set_title(perturb)
        ax.set_ylabel("PR-AUC")
        style_axes(ax)
    axes[1, 1].legend(frameon=False, loc="upper right")
    fig.suptitle("Robustness under controlled test-time perturbations", fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    save_fig(fig, "figure_5_robustness.png")

    # Figure 6: text models and fusion comparison.
    fig, ax = plt.subplots(figsize=(7.6, 4.5))
    for model in ["lr_tfidf", "lr_minilm", "fusion_tfidf", "fusion_minilm"]:
        g = exp1[exp1["Model"] == model].sort_values("Fraction")
        ax.plot(g["Fraction"] * 100, g["PR-AUC"], marker="o", lw=1.8, label=model, color=color_map[model])
    ax.set_xlabel("Labeled training data (%)")
    ax.set_ylabel("PR-AUC")
    ax.set_title("Classical TF-IDF, frozen MiniLM, and fusion baselines", fontweight="bold")
    style_axes(ax)
    ax.legend(frameon=False)
    save_fig(fig, "figure_6_text_comparison.png")

    # Figure 7: calibration summary.
    perf = read_csv("table_6_predictive_performance.csv")
    fig, ax = plt.subplots(figsize=(8.5, 4.6))
    x = np.arange(len(models))
    ax.bar(x - 0.18, [perf.loc[perf.Model == m, "Brier Score"].iloc[0] for m in models], width=0.36, color=COLORS["teal"], label="Brier score")
    ax.bar(x + 0.18, [perf.loc[perf.Model == m, "ECE"].iloc[0] for m in models], width=0.36, color=COLORS["amber"], label="ECE")
    ax.set_xticks(x)
    ax.set_xticklabels([MODEL_LABELS[m] for m in models])
    ax.set_ylabel("Calibration error")
    ax.set_title("Calibration summary after validation-set sigmoid scaling", fontweight="bold")
    style_axes(ax)
    ax.legend(frameon=False)
    save_fig(fig, "figure_7_calibration_summary.png")

    # Figure 8: routing distribution.
    routing = read_csv("table_10_routing_results.csv")
    fig, ax = plt.subplots(figsize=(8.6, 4.8))
    auto = np.array([routing.loc[routing.Model == m, "Auto-Clear Rate"].iloc[0] for m in models])
    review = np.array([routing.loc[routing.Model == m, "Human-Review Rate"].iloc[0] for m in models])
    urgent = np.array([routing.loc[routing.Model == m, "Urgent-Inspection Rate"].iloc[0] for m in models])
    ax.bar(x, auto * 100, color=COLORS["green"], label="Auto-clear")
    ax.bar(x, review * 100, bottom=auto * 100, color=COLORS["amber"], label="Human review")
    ax.bar(x, urgent * 100, bottom=(auto + review) * 100, color=COLORS["red"], label="Urgent inspection")
    ax.set_xticks(x)
    ax.set_xticklabels([MODEL_LABELS[m] for m in models])
    ax.set_ylabel("Test cases (%)")
    ax.set_ylim(0, 100)
    ax.set_title("Validation-selected decision routing on the test set", fontweight="bold")
    style_axes(ax)
    ax.legend(frameon=False, ncol=3, loc="upper center", bbox_to_anchor=(0.5, 1.12))
    save_fig(fig, "figure_8_routing_distribution.png")

    # Figure 9: routing sensitivity.
    sens = read_csv("routing_sensitivity.csv")
    fig, ax = plt.subplots(figsize=(7.8, 4.6))
    for model in ["lr_struct", "rf_struct", "xgb_struct", "fusion_minilm"]:
        g = sens[sens["Model"] == model].sort_values("alpha")
        ax.plot(g["alpha"], g["test_workload_reduction"] * 100, marker="o", lw=2, label=model, color=color_map[model])
    ax.set_xlabel(f"Validation missed-fault ceiling ({ALPHA})")
    ax.set_ylabel("Test workload reduction (%)")
    ax.set_title("Workload reduction sensitivity to safety constraint", fontweight="bold")
    style_axes(ax)
    ax.legend(frameon=False)
    save_fig(fig, "figure_9_sensitivity.png")

    # Figure 10: global feature importance.
    imp = read_csv("global_feature_importance.csv").head(12).iloc[::-1]
    fig, ax = plt.subplots(figsize=(7.2, 5.2))
    ax.barh(imp["Feature"], imp["Importance"], color=COLORS["blue"])
    ax.set_xlabel("Mean absolute SHAP value")
    ax.set_title("Global feature importance for structured XGBoost", fontweight="bold")
    style_axes(ax)
    save_fig(fig, "figure_10_shap_importance.png")

    # Figure 11: local SHAP compact diagnostic.
    local_path = OUT / "local_shap_contributions.csv"
    if local_path.exists():
        local = pd.read_csv(local_path)
        fig, axes = plt.subplots(2, 2, figsize=(10, 7))
        for ax, case in zip(axes.ravel(), ["TP", "FP", "FN", "TN"]):
            g = local[local["Case"] == case].sort_values("Abs SHAP Contribution", ascending=True).tail(8)
            colors = [COLORS["red"] if v > 0 else COLORS["blue"] for v in g["SHAP Contribution"]]
            ax.barh(g["Feature"], g["SHAP Contribution"], color=colors)
            ax.axvline(0, color="#111827", lw=0.8)
            prob = g["Predicted Probability"].iloc[0] if len(g) else np.nan
            ax.set_title(f"{case}: p={prob:.3f}" if len(g) else case)
            style_axes(ax)
        fig.suptitle("Local SHAP diagnostics for four structured XGBoost cases", fontweight="bold")
        fig.tight_layout(rect=[0, 0, 1, 0.95])
        save_fig(fig, "figure_11_local_shap.png")

    make_graphical_abstract()


def make_graphical_abstract():
    path = FIG / "GRAPHICAL_ABSTRACT_MDPI_AI.png"
    w, h = 1600, 820
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 44)
        head_font = ImageFont.truetype("arial.ttf", 28)
        body_font = ImageFont.truetype("arial.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        title_font = head_font = body_font = small_font = ImageFont.load_default()

    draw.rectangle([0, 0, w, 92], fill="#EEF6FF")
    draw.text((54, 25), "Synthetic Low-Data Industrial O&M Benchmark", fill=COLORS["ink"], font=title_font)
    draw.text((56, 104), "Question", fill=COLORS["blue"], font=head_font)
    draw.text((56, 142), "Can AI maintenance decision support remain reliable with scarce labels, rare faults,\nnoisy sensors, imperfect reports, and deployment-site shift?", fill=COLORS["ink"], font=body_font)

    boxes = [
        ("Synthetic generator", "six plants\n480 assets\n175,200 records", 72, 300, COLORS["blue"]),
        ("Model stack", "structured\nTF-IDF\nfrozen MiniLM\nfusion", 428, 300, COLORS["teal"]),
        ("Trust tests", "low labels\nheld-out sites\nnoise/missingness\ncalibration", 784, 300, COLORS["purple"]),
        ("Decision routing", "validation thresholds\nmissed-fault ceiling\nworkload reduction", 1140, 300, COLORS["green"]),
    ]
    for title, body, x0, y0, color in boxes:
        draw.rounded_rectangle([x0, y0, x0 + 300, y0 + 230], radius=28, outline=color, width=5, fill="white")
        draw.text((x0 + 26, y0 + 25), title, fill=color, font=head_font)
        draw.text((x0 + 26, y0 + 78), body, fill=COLORS["ink"], font=body_font, spacing=8)
    for x0 in [374, 730, 1086]:
        draw.line([x0, 415, x0 + 54, 415], fill=COLORS["muted"], width=5)
        draw.polygon([(x0 + 54, 415), (x0 + 36, 405), (x0 + 36, 425)], fill=COLORS["muted"])

    draw.rectangle([54, 605, 1546, 736], fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw.text((82, 628), "Main synthetic-result pattern:", fill=COLORS["ink"], font=head_font)
    draw.text((82, 675), "structured models dominate text-only baselines; random splits overstate site generalization;\nvalidation-selected routing exposes when auto-clear is unsafe.", fill=COLORS["ink"], font=body_font, spacing=6)
    draw.text((56, 770), "Fully synthetic data only; no real industrial deployment validation is claimed.", fill=COLORS["muted"], font=small_font)
    img.save(path, "PNG")


def build_docx():
    SUB.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)

    metadata = json.loads((OUT / "generation_metadata.json").read_text(encoding="utf-8"))
    perf = read_csv("table_6_predictive_performance.csv")
    site = read_csv("table_7_cross_site_generalization.csv")
    routing_thr = read_csv("table_9_routing_thresholds.csv")
    routing = read_csv("table_10_routing_results.csv")
    imp = read_csv("global_feature_importance.csv").head(10)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("A Synthetic Low-Data Benchmark for Trustworthy AI-Based Industrial Operation and Maintenance Decision Support")
    p = doc.add_paragraph("Article Type: Article")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Author: Parth Purohit")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Affiliation: Independent Researcher")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Correspondence: To be supplied in the MDPI submission system.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "Keywords: predictive maintenance; industrial operation and maintenance; synthetic benchmark; low-data learning; domain shift; uncertainty calibration; human-in-the-loop decision support; technician reports; explainable AI; reproducibility")

    add_paragraph(doc, "Abstract", "Heading 1")
    add_paragraph(doc, "Industrial operation and maintenance (O&M) AI is difficult to evaluate when labeled failures are scarce, observations are noisy, technician reports are incomplete, and deployment sites differ from training sites. This article presents a reproducible synthetic benchmark for these conditions in a generic multi-site industrial setting. The generator simulates 480 assets across six plants over 365 days, producing 175,200 shift-level records with hidden degradation, noisy structured observations, imperfect technician reports, leakage-audited features, chronological splits, random splits, and site-held-out splits. Seven calibrated baselines are evaluated: structured-only classifiers, TF-IDF plus Logistic Regression, frozen MiniLM sentence embeddings, and structured-plus-text fusion. Random Forest obtains the highest chronological-test PR-AUC (0.398), but at the default 0.5 threshold still misses 58.5% of positives. Random-split evaluation overstates generalization: Random Forest PR-AUC falls from 0.708 under random splitting to 0.439 under held-out-site testing, and false negative rate increases from 0.471 to 0.815. Validation-set routing thresholds can reduce workload, but only for models whose low-risk calibration tails satisfy the missed-fault constraint. The study is fully synthetic and does not claim real industrial deployment validation.")

    add_paragraph(doc, "1. Introduction", "Heading 1")
    add_paragraph(doc, "Industrial O&M decision support is usually evaluated under constraints that are difficult to reproduce: rare failure labels, heterogeneous operating sites, noisy condition-monitoring signals, missing observations, and unstructured technician notes. Accuracy or ROC-AUC alone is therefore not enough. A model can rank cases well and still be unsafe at an operational threshold if false negatives are high or probabilities are poorly calibrated.")
    add_paragraph(doc, "The limited-data problem is not a secondary implementation detail in predictive maintenance. Failure events are often the labels of greatest operational interest, but they are also the least frequently observed. When a system has many normal observations and very few positive maintenance events, a model can appear strong under conventional accuracy while providing little useful support for the cases that matter most. The local literature review for this study identifies this pattern across work on scarce failures, class imbalance, noisy observations, and rare-event analysis in manufacturing and condition monitoring [1-6]. This motivates an evaluation design that reports PR-AUC, recall, false negative rate, calibration, and routing behavior rather than relying on a single aggregate score.")
    add_paragraph(doc, "Another practical issue is that industrial deployment rarely matches the clean random-split assumption. A model trained on one machine, condition, or plant may be evaluated in a different site with different usage intensity, maintenance timing, environmental conditions, sensor calibration, or report-writing conventions. The domain-adaptation and benchmarking literature shows that performance can change materially when evaluation uses unseen domains or more conservative split procedures [13-17,31]. For this reason, the benchmark is designed around both chronological testing and held-out-site testing. The goal is not to claim that the synthetic plants represent real facilities, but to provide a controlled protocol in which site shift is explicit and measurable.")
    add_paragraph(doc, "The third motivation is multimodality. Structured sensor and maintenance-history variables are common in predictive-maintenance studies, but maintenance supervisors also use short notes, work orders, and inspection comments. Prior studies show that free-form maintenance text can contain predictive information, while also being noisy, abbreviated, technical, and difficult for generic language models [18-22]. This paper therefore compares structured-only, text-only, and structured-plus-text baselines. The text component is deliberately conservative: TF-IDF plus Logistic Regression is included as a transparent sparse baseline, and MiniLM is used only as a frozen general-purpose sentence encoder. The benchmark does not assume that modern embeddings should outperform simple text baselines.")
    add_paragraph(doc, "This work contributes a controlled synthetic benchmark and evaluation protocol for low-data industrial O&M. The benchmark is intentionally generic: it contains plants, assets, components, sensor observations, technician reports, maintenance events, and human maintenance-supervisor routing decisions. It does not use or represent any real industrial, personal, proprietary, institutional, or project-specific operational data.")
    add_paragraph(doc, "The benchmark contribution is methodological rather than model-centric. It contributes a generator, leakage checks, fixed split files, label-scarcity splits, site-held-out splits, robustness scenarios, calibrated prediction pipelines, threshold selection on validation data, and publication artifacts. These design choices follow the broader reproducibility argument that benchmarks need stable data generation, stable partitions, visible configuration, and repeatable evaluation outputs [31-33]. A single high-performing model on one synthetic dataset would be a weak contribution; a reusable protocol that exposes failure modes under scarcity, noise, shift, calibration error, and constrained routing is the intended contribution.")
    add_paragraph(doc, "The research questions are: RQ1, how do structured, text-based, and fusion models behave as labels become scarce? RQ2, how much do performance and calibration degrade under held-out-site testing? RQ3, how robust are models to sensor noise, missing structured values, and ambiguous or missing technician reports? RQ4, does a frozen MiniLM sentence encoder improve technician-report modeling compared with TF-IDF? RQ5, can validation-set-optimized routing reduce workload while constraining missed faults? RQ6, are the resulting explanations plausible diagnostic artifacts rather than leakage artifacts?")

    add_paragraph(doc, "2. Related Work", "Heading 1")
    add_paragraph(doc, "2.1. Low-Data Predictive Maintenance and Rare Events", "Heading 2")
    add_paragraph(doc, "The first literature stream concerns data scarcity, imbalance, and noisy fault labels. Hakami [1] is especially relevant because it frames predictive maintenance around scarcity, imbalance, feature selection, and temporal learning. The source notes for this manuscript record that the study used a production-plant condition-monitoring dataset with a severe healthy-versus-failure imbalance. Ramirez-Sanz et al. [2] review semi-supervised industrial fault detection and diagnosis, which is directly relevant to the limited-label premise even though this manuscript does not implement a semi-supervised method. Jalayer et al. [3], Kafunah et al. [4], Swana et al. [5], and Shyalika et al. [6] all reinforce the broader point that industrial fault recognition often requires evaluation choices designed for imbalanced, noisy, and rare-event settings. These papers justify the use of PR-AUC, macro-F1, recall, and false negative rate as primary outputs.")
    add_paragraph(doc, "The present work differs from that stream in two ways. First, it does not propose another imbalance-handling algorithm such as SMOTE, Tomek-link processing, data enrichment, or a new deep architecture. Second, it does not claim that a synthetic generator can solve the lack of real failures. Instead, it creates a controllable benchmark in which failure rarity, label scarcity, missingness, noise, and site shift can be changed deliberately. This is a narrower and more auditable contribution: the benchmark can reveal which model families become unstable as labels disappear, but it remains a simulation and cannot replace real operational validation.")
    add_paragraph(doc, "2.2. Synthetic Data, Digital Twins, and Benchmark Construction", "Heading 2")
    add_paragraph(doc, "The second literature stream concerns digital twins, simulation, and synthetic data for predictive maintenance. Chen et al. [7] describe the growing role of machine learning in digital-twin predictive maintenance, including fault diagnosis, health indicators, and remaining-useful-life prediction. Hosamo et al. [8] and Singh et al. [9] provide examples of digital-twin predictive-maintenance frameworks, while Yan et al. [10] directly links simulated fault data to imbalanced fault diagnosis through a digital-twin-assisted framework. Mikolajewska et al. [11] discuss generative AI in AI-based digital twins, and Moccardi et al. [12] contributes to the uncertainty-aware predictive-maintenance literature through a conformal framework. Together, these sources support synthetic and simulation-based evaluation as useful research tools, provided that their validation boundaries are explicit.")
    add_paragraph(doc, "This boundary is important. The generator in this paper should be read as a benchmark simulator, not as a physical digital twin of a pump, compressor, conveyor, or motor. Its latent degradation process is designed to create plausible statistical relationships between age, load, maintenance history, noisy sensors, reports, and labels. It does not solve the harder engineering problem of fitting physical degradation parameters to a real machine. The paper therefore reports synthetic evidence only and uses the digital-twin literature primarily to motivate controlled simulation and to state the limitations of simulation-to-real transfer.")
    add_paragraph(doc, "2.3. Domain Shift and Evaluation Splits", "Heading 2")
    add_paragraph(doc, "Domain shift is central to the benchmark. Chen et al. [13] study fault diagnosis under unseen conditions using adversarial domain-invariant generalization. Asutkar and Tallur [14] and Ragab et al. [15] similarly address transfer or domain adaptation in machine fault diagnosis and remaining-useful-life prediction. Forest and Fink [16] connect domain adaptation with calibration, showing that confidence quality can matter when target-domain pseudo-labels are used. Wen et al. [17] further motivates domain adaptation for remaining useful life prediction when source and target distributions differ. These papers support the held-out-site experiment, where the test plant is deliberately not present in training.")
    add_paragraph(doc, "The split-design issue is also methodological. Hendriks et al. [31] show that benchmark conclusions can be distorted when common split procedures allow overlap that is not realistic for deployment. Their work is specific to a bearing fault dataset, but the principle generalizes: if a benchmark asks a deployment question, the split must represent that question. This manuscript therefore reports both random-split and site-held-out performance. The random split is useful as a familiar baseline, but the held-out-site protocol is more aligned with transfer to a new industrial location.")
    add_paragraph(doc, "2.4. Technician Reports and Maintenance Text", "Heading 2")
    add_paragraph(doc, "The third literature stream concerns maintenance text. Usuga-Cadavid et al. [18] show that free-form maintenance logs can be valuable for predictive-maintenance tasks, but also emphasize the need for interpretability and the difficulty of unstructured maintenance language. Sundaram and Zeid [19] focus on technical language processing for prognostics and health management work orders, which supports the claim that work-order text can contain specialized shorthand and domain-specific vocabulary. Giordano and Fantoni [20] use NLP to decompose maintenance actions into subtasks, Pavlopoulos et al. [21] combine ML and NLP for fault nowcasting, and Zhou et al. [22] provide a recent multimodal time-series/text fault-diagnosis framework.")
    add_paragraph(doc, "These studies justify including technician reports in the benchmark, but they also caution against overclaiming. The reports generated here are template-based, intentionally sparse, and sometimes ambiguous or misleading. They are not equivalent to real technician language. For that reason, the paper treats text results as a stress test of modality value under controlled report ambiguity rather than as evidence about real work-order NLP. This distinction becomes important in the results: text-only models are weak in this benchmark, but that does not imply that real maintenance text is unhelpful. It implies only that the synthetic text channel, as designed, does not dominate structured condition-monitoring features.")
    add_paragraph(doc, "2.5. Calibration, Reject Options, and Human Review", "Heading 2")
    add_paragraph(doc, "Calibration and abstention are the fourth literature stream. Silva Filho et al. [23] survey how to assess and improve predicted class probabilities, including reliability diagrams, proper scoring, and post-hoc calibration. Hendrickx et al. [24] review machine learning with a reject option, which provides a general basis for routing uncertain predictions to human review. Hasan et al. [25] discuss reject-option and post-training processing for trustworthy neural networks, and Sayin et al. [26] argue that classifiers in critical cyber-physical systems should be designed with rejection, calibration, and cost-sensitive thresholds in mind. Perez-Cerrolaza et al. [27] further situates AI within safety-critical industrial systems.")
    add_paragraph(doc, "The routing experiment in this manuscript is therefore grounded in the reject-option literature, but it remains a benchmark-specific formulation. The paper uses three categories: auto-clear, human review, and urgent inspection. It then selects t_low and t_high on the validation set to maximize workload reduction subject to a missed-fault ceiling. The alpha values are not presented as industrial safety standards. They are sensitivity-analysis settings that expose how much workload reduction remains possible when the benchmark imposes stricter or looser safety constraints.")
    add_paragraph(doc, "2.6. Explainability and Reproducibility", "Heading 2")
    add_paragraph(doc, "Explainability is included as a diagnostic layer rather than a claim of causal understanding. Cummins et al. [28] survey explainable predictive maintenance and show that methods such as SHAP and LIME are widely used but not universally validated as explanation-quality standards. Brusa et al. [29] and Gawde et al. [30] provide additional evidence that feature-contribution tools are common in industrial condition monitoring and rotating-machine predictive maintenance. This paper uses SHAP for tree-based structured models and coefficient inspection for sparse text models, while avoiding direct interpretation of individual MiniLM embedding dimensions.")
    add_paragraph(doc, "Finally, reproducibility is treated as part of the contribution. Barnard [32] argues for turning scientific datasets into reproducible benchmarks, and McDermott et al. [33] highlight continuing reproducibility problems in machine-learning research. In response, this benchmark releases code, configuration, random seeds, generated data, split files, metrics, tables, and figures. The local literature review folder is used only as an evidence base for manuscript writing; downloaded third-party PDFs and Consensus reports are not part of the public reproducibility package.")

    add_paragraph(doc, "3. Materials and Methods", "Heading 1")
    add_paragraph(doc, "3.1. Synthetic Benchmark Design", "Heading 2")
    add_paragraph(doc, f"The dataset contains {metadata['row_count']:,} records, {metadata['asset_count']} assets, six plants, five asset types, and daily or shift-level observations from {metadata['date_min']} to {metadata['date_max']}. The main positive rate is {100 * metadata['maintenance_required_7d_positive_rate']:.2f}%, and technician reports are present for {100 * metadata['technician_report_coverage']:.1f}% of records.")
    add_caption(doc, "Table 1. Synthetic dataset summary.")
    make_table(doc, ["Element", "Value"], [
        ["Plants", "6 simulated industrial sites"],
        ["Assets", "480 total assets; 80 per plant"],
        ["Records", f"{metadata['row_count']:,} daily/shift-level rows"],
        ["Main label", "maintenance_required_7d"],
        ["Positive rate", f"{100 * metadata['maintenance_required_7d_positive_rate']:.2f}%"],
        ["Technician report coverage", f"{100 * metadata['technician_report_coverage']:.1f}%"],
        ["Data source", "Fully synthetic; no real operational records"],
    ], widths=[2.2, 4.4])
    add_paragraph(doc, "Table 1 gives the scale and boundary of the generated benchmark. The dataset is large enough to train conventional tabular and text models, but the positive class remains intentionally rare. This is the key experimental tension: the total row count is not small, yet the number of safety-relevant positive examples is limited relative to normal operation. The technician-report coverage is also intentionally incomplete. A model cannot rely on text being present for every case, and the absence of a report is itself part of the low-information operating condition.")

    add_paragraph(doc, "The hidden degradation state is denoted by D_t. It is not included as a model input. A simplified daily degradation update is:")
    add_equation(doc, f"D_t = D_(t-1) + {DELTA}d_base {DOT} {GAMMA}_site {DOT} (1 + {BETA} {DOT} load) + {ETA}_t + {THETA}_t")
    add_paragraph(doc, f"where {DELTA}d_base is the baseline wear rate, {GAMMA}_site is a site-specific scaling factor, {BETA} represents sensitivity to load, {ETA}_t ~ N(0, {SIGMA}{SUP2}) is continuous drift noise, and {THETA}_t represents random shocks. When D_t {GE} 0.75, a physical fault state can be triggered. Maintenance resets D_t to a low post-maintenance range. This formulation is a benchmark simulator, not a physically validated digital twin.")
    add_caption(doc, "Figure 1. Synthetic benchmark and evaluation flow.")
    doc.add_picture(str(FIG / "figure_1_pipeline.png"), width=Inches(6.6))
    add_caption(doc, "Figure 2. Hidden degradation and noisy observable sensor response for a representative asset.")
    doc.add_picture(str(FIG / "figure_2_degradation.png"), width=Inches(6.4))
    add_paragraph(doc, "Figures 1 and 2 clarify the role of the latent degradation variable. The simulator first evolves hidden degradation and then generates sensors, reports, maintenance events, and labels from that hidden process. The hidden state is shown only for benchmark explanation and debugging. It is never provided as a model input. This separation is central to the leakage boundary: models must infer future maintenance requirement from noisy observable consequences rather than from the internal variable that helped generate the label.")

    add_caption(doc, "Table 2. Feature groups used in the benchmark.")
    make_table(doc, ["Group", "Columns"], [
        ["Asset metadata", "date, asset_id, site_id, asset_type, component_type"],
        ["Operational state", "age, operating_hours, load_factor, duty_cycle, shift_type"],
        ["Sensors", "ambient_temp, humidity, vibration_rms, vibration_kurtosis, acoustic_level, motor_current, temperature, pressure, and flow deviations"],
        ["Maintenance history", "time_since_last_maintenance, previous_fault_count"],
        ["Text", "technician_report"],
        ["Labels", "maintenance_required_7d, fault_type, maintenance_priority, remaining_useful_life_bin"],
    ], widths=[1.55, 5.2], font_size=8)
    add_paragraph(doc, "Table 2 separates the benchmark into feature groups because the experimental comparisons are modality-based. Structured-only models receive metadata, operational state, sensors, and maintenance history. Text-only models receive technician reports. Fusion models receive both. Labels are stored with the generated dataset for evaluation, but only the binary maintenance_required_7d label is used as the main prediction target in the experiments reported here. The auxiliary labels are retained to support future multiclass or ordinal extensions.")

    add_paragraph(doc, "Technician reports are generated from templates conditioned on fault state, sensor anomalies, maintenance state, and random ambiguity. A generated report has a 15% probability of being drawn from a mismatched template, and 10% of generated reports are replaced by vague entries. Reports are sparse, with probability increasing near high degradation or maintenance events.")
    add_paragraph(doc, "A leakage audit fails loudly if forbidden columns such as latent_degradation, health_score, true_risk_score, future_fault_count, future_maintenance_flag, post_event_status, maintenance_done_after_prediction, or remaining_useful_life_true are present in the training feature matrix.")
    add_paragraph(doc, "The report-generation design is intentionally imperfect. In real maintenance settings, text can be incomplete, ambiguous, delayed, or written after a human has already formed an operational hypothesis. The benchmark approximates this by making some reports vague or mismatched and by leaving many records without reports. This design choice prevents the text channel from becoming an unrealistically clean label proxy. It also creates a useful negative-control question: if a modern sentence encoder cannot outperform a sparse baseline on these short reports, the result should be interpreted as evidence about this text channel, not as evidence against maintenance-text mining generally.")

    add_paragraph(doc, "3.2. Experimental Protocol", "Heading 2")
    add_caption(doc, "Table 3. Experimental scenarios and controlled variables.")
    make_table(doc, ["Experiment", "Controlled variable", "Main outputs"], [
        ["Low-label learning", "1%, 5%, 10%, 25%, 50%, 100% labeled training fractions", "PR-AUC, ROC-AUC, Brier score, ECE"],
        ["Cross-site shift", "Six rotating held-out plants", "Random-vs-site PR-AUC gap, FNR inflation, ECE"],
        ["Robustness", "Sensor noise, missing sensors, missing reports, report ambiguity", "PR-AUC, macro-F1, recall, FNR"],
        ["Decision routing", f"Validation-selected t_low and t_high under {ALPHA} = 0.05 plus sensitivity", "Auto-clear, human review, urgent inspection, missed-fault rate"],
    ], widths=[1.45, 2.8, 2.65], font_size=8)
    add_paragraph(doc, "Table 3 organizes the experiments by the failure mode they are intended to expose. The low-label experiment tests whether a method is data-efficient. The cross-site experiment tests whether random-split conclusions survive deployment-style shift. The robustness experiment separates sensor noise, sensor missingness, report missingness, and report ambiguity, because these degradations can affect modalities differently. The routing experiment translates calibrated probabilities into operational categories, making explicit the trade-off between reducing supervisor workload and avoiding missed maintenance positives.")

    add_paragraph(doc, "The default chronological split uses Plants 1 to 5 only: 70% of dates for training, 15% for validation, and 15% for testing. The random split uses all plants with stratification by the binary label. The site-held-out experiment rotates each of the six plants as an unseen test site.")
    doc.add_page_break()
    add_paragraph(doc, "3.3. Models, Calibration, and Routing", "Heading 2")
    add_caption(doc, "Table 4. Model stack and modalities.")
    make_table(doc, ["Model ID", "Model", "Input"], [
        ["lr_struct", "Logistic Regression", "Structured features"],
        ["rf_struct", "Random Forest", "Structured features"],
        ["xgb_struct", "XGBoost", "Structured features"],
        ["lr_tfidf", "TF-IDF + Logistic Regression", "Technician reports"],
        ["lr_minilm", "Frozen all-MiniLM-L6-v2 + Logistic Regression", "Technician reports"],
        ["fusion_tfidf", "XGBoost", "Structured + TF-IDF"],
        ["fusion_minilm", "XGBoost", "Structured + frozen MiniLM"],
    ], widths=[1.1, 2.6, 2.8], font_size=8)
    add_paragraph(doc, "Table 4 deliberately mixes simple and stronger baselines. Logistic Regression is included as a transparent linear baseline, Random Forest and XGBoost represent widely used nonlinear tabular learners, TF-IDF plus Logistic Regression gives a sparse and interpretable text baseline, and the frozen MiniLM baselines test whether a compact dense sentence representation helps with short technician reports. MiniLM is not fine-tuned in the primary experiments because full fine-tuning would add variance and training complexity in exactly the low-label setting that the benchmark is meant to study.")
    add_paragraph(doc, "TF-IDF features are extracted from technician reports, and MiniLM text representations use the frozen all-MiniLM-L6-v2 encoder. The choice of a frozen encoder is methodological, not a claim that MiniLM is an industrial maintenance language model. Tree and linear baselines are calibrated on the validation set using sigmoid calibration, following the general principle that predicted probabilities should be assessed and calibrated before they are used in decision support [23,26].")
    add_paragraph(doc, f"The routing policy uses calibrated probability p = P(y = 1 | x). If p {LE} t_low, the case is auto-cleared; if t_low < p < t_high, it is routed to human review; if p {GE} t_high, it is routed to urgent inspection. Thresholds are chosen on the validation set only:")
    add_equation(doc, f"(t_low*, t_high*) = arg max WorkloadReduction(t_low, t_high), subject to MCFR {LE} {ALPHA}")
    add_equation(doc, f"MCFR = {SUM} 1[y_i = 1 {WEDGE} p_i {LE} t_low] / {SUM} 1[y_i = 1]")
    add_paragraph(doc, "The selected thresholds are frozen before final test evaluation.")
    add_paragraph(doc, "The routing equations formalize the reject-option idea in an O&M vocabulary. Auto-clear corresponds to a low-risk automated decision, human review corresponds to abstention or deferral, and urgent inspection corresponds to a high-risk action recommendation. The validation-only threshold search is essential. If thresholds were chosen after viewing test results, the routing layer would be test-set tuned and the reported missed-fault rate would be optimistic. The benchmark therefore treats threshold selection as part of model development and applies the selected pair unchanged to the held-out test set.")

    add_paragraph(doc, "4. Results", "Heading 1")
    add_paragraph(doc, "4.1. Label Scarcity and Main Predictive Performance", "Heading 2")
    add_caption(doc, "Figure 3. Label-scarcity learning curves.")
    doc.add_picture(str(FIG / "figure_3_learning_curves.png"), width=Inches(6.5))
    add_caption(doc, "Table 5. Main chronological-test performance at 100% labeled training data.")
    rows = []
    for _, r in perf.iterrows():
        rows.append([r["Model"], fmt(r["ROC-AUC"]), fmt(r["PR-AUC"]), fmt(r["Macro-F1"]), fmt(r["Recall (Sensitivity)"]), fmt(r["False Negative Rate"]), fmt(r["Brier Score"]), fmt(r["ECE"])])
    make_table(doc, ["Model", "ROC-AUC", "PR-AUC", "Macro-F1", "Recall", "FNR", "Brier", "ECE"], rows, font_size=7.2)
    add_paragraph(doc, "Random Forest is the strongest chronological-test baseline by PR-AUC (0.398), followed by MiniLM fusion and XGBoost. The text-only baselines have weak ranking performance and predict no positives at the default 0.5 threshold, confirming that the short synthetic reports alone are insufficient for reliable decision support.")
    add_paragraph(doc, "Figure 3 shows that the structured models are relatively stable once a small fraction of labels is available, whereas the text-only curves remain close to the rare-event baseline. This pattern is consistent with the generator design: the structured variables contain noisy but direct consequences of degradation, while the reports are sparse and deliberately ambiguous. The learning-curve result does not prove that low-label learning is solved. It shows that, in this synthetic configuration, the structured signal is strong enough that Random Forest and XGBoost can rank positives better than text-only models even when labels are restricted.")
    add_paragraph(doc, "Table 5 shows why PR-AUC and FNR are both needed. Random Forest has the highest PR-AUC, but its recall at the default threshold is only 0.415 and its false negative rate is 0.585. Logistic Regression has lower PR-AUC but slightly higher recall at the same threshold. These differences would be obscured by a single ranking metric. For maintenance decision support, the default threshold is rarely the final operating point, but reporting it is still useful because it reveals whether a model is naturally conservative or permissive before routing thresholds are optimized.")

    add_paragraph(doc, "4.2. Cross-Site Generalization", "Heading 2")
    add_caption(doc, "Figure 4. Random split versus held-out-site PR-AUC.")
    doc.add_picture(str(FIG / "figure_4_site_shift.png"), width=Inches(6.5))
    add_caption(doc, "Table 6. Random-split versus held-out-site generalization.")
    rows = []
    for _, r in site.iterrows():
        rows.append([r["Model"], fmt(r["Random-Split Test PR-AUC"]), fmt(r["Cross-Site Mean PR-AUC"]), fmt(r["PR-AUC Drop (Shift Gap)"]), fmt(r["Random-Split FNR"]), fmt(r["Cross-Site Mean FNR"]), fmt(r["FNR Inflation"])])
    make_table(doc, ["Model", "Random PR-AUC", "Site PR-AUC", "PR-AUC drop", "Random FNR", "Site FNR", "FNR inflation"], rows, font_size=7.0)
    add_paragraph(doc, "Random splitting substantially overstates deployment-style performance for tree and fusion models. For Random Forest, mean held-out-site PR-AUC is 0.439 compared with 0.708 under random splitting, while mean FNR rises to 0.815.")
    add_paragraph(doc, "Figure 4 and Table 6 are the strongest evidence that the benchmark is doing more than measuring generic classification accuracy. Under the random split, the model sees records from all sites during training and testing. Under held-out-site testing, the test plant is unseen. The resulting gap is large for Random Forest and the fusion models, which suggests that these models exploit site-specific patterns that help under random mixing but do not transfer cleanly to a new plant. The result supports the decision to report both split styles: random-split results are not false, but they answer an easier question.")
    add_paragraph(doc, "The false negative inflation under site shift is operationally important. A model that misses a higher fraction of positives on an unseen plant would be risky if deployed as an automatic maintenance filter. The result also explains why calibration and routing are evaluated separately. A model can rank cases well under one split and still become poorly suited for auto-clear decisions under another distribution. This is aligned with the literature on unseen operating conditions and benchmark split design [13-17,31].")

    add_paragraph(doc, "4.3. Robustness and Calibration", "Heading 2")
    add_caption(doc, "Figure 5. Robustness under controlled test-time perturbations.")
    doc.add_picture(str(FIG / "figure_5_robustness.png"), width=Inches(6.5))
    add_caption(doc, "Figure 6. Classical TF-IDF, frozen MiniLM, and fusion baselines under label scarcity.")
    doc.add_picture(str(FIG / "figure_6_text_comparison.png"), width=Inches(6.1))
    add_caption(doc, "Figure 7. Calibration summary after validation-set sigmoid scaling.")
    doc.add_picture(str(FIG / "figure_7_calibration_summary.png"), width=Inches(6.4))
    robust = read_csv("exp3_robustness_results.csv")
    selected = robust[(robust["Model"].isin(["rf_struct", "fusion_minilm", "lr_tfidf"])) & (
        ((robust["Perturbation_Type"] == "Sensor Noise") & (robust["Level"].isin(["1.0x", "3.0x"]))) |
        ((robust["Perturbation_Type"] == "Missing Sensors") & (robust["Level"].isin(["40%"]))) |
        ((robust["Perturbation_Type"] == "Missing Reports") & (robust["Level"].isin(["75%"]))) |
        ((robust["Perturbation_Type"] == "Report Ambiguity") & (robust["Level"].isin(["75%"])))
    )]
    rows = [[r["Model"], r["Perturbation_Type"], r["Level"], fmt(r["PR-AUC"]), fmt(r["Macro-F1"]), fmt(r["Recall"]), fmt(r["FNR"])] for _, r in selected.iterrows()]
    add_caption(doc, "Table 7. Selected robustness results.")
    make_table(doc, ["Model", "Perturbation", "Level", "PR-AUC", "Macro-F1", "Recall", "FNR"], rows, font_size=7.0)
    add_paragraph(doc, "Figure 5 separates sensor degradation from report degradation. Sensor noise reduces PR-AUC for the structured and fusion models, which is expected because these models rely heavily on vibration, temperature, acoustic, pressure, and maintenance-history variables. Report missingness and report ambiguity have a much smaller effect on Random Forest because it does not use text, and a limited effect on the fusion model because structured features dominate. The text-only TF-IDF baseline remains weak across these perturbations, which again indicates that the generated reports are not sufficiently informative on their own.")
    add_paragraph(doc, "Figure 6 compares sparse TF-IDF and frozen MiniLM representations under label scarcity. MiniLM does not provide a meaningful text-only improvement over TF-IDF, and fusion with MiniLM does not surpass the best structured-only model. This is a useful negative result. It prevents the manuscript from implying that a modern sentence encoder automatically improves maintenance decision support. In this benchmark, the main value of text is limited by report sparsity and ambiguity, not only by the representation model.")
    add_paragraph(doc, "Figure 7 and Table 7 also show that calibration and robustness should not be collapsed into one statement. The Brier and ECE values are moderate after validation-set sigmoid calibration, but calibration summaries alone do not reveal the severe effect of test-time sensor noise on PR-AUC. Conversely, a model may preserve ranking under missing reports because the report channel is weak, while still being poorly calibrated for routing. These results justify the manuscript's multi-metric approach.")

    add_paragraph(doc, "4.4. Validation-Selected Decision Routing", "Heading 2")
    add_caption(doc, "Figure 8. Test-set decision-routing distribution by model.")
    doc.add_picture(str(FIG / "figure_8_routing_distribution.png"), width=Inches(6.5))
    add_caption(doc, "Figure 9. Workload reduction sensitivity to validation missed-fault ceiling.")
    doc.add_picture(str(FIG / "figure_9_sensitivity.png"), width=Inches(6.1))
    doc.add_page_break()
    merged = routing_thr.merge(routing, on="Model")
    rows = []
    for _, r in merged.iterrows():
        rows.append([r["Model"], fmt(r["t_low*"], 2), fmt(r["t_high*"], 2), pct(r["Validation Workload Reduction"], 1), pct(r["Validation Missed Fault Rate"], 2), pct(r["Workload Reduction"], 1), pct(r["Missed Critical Fault Rate"], 2), pct(r["Urgent-Inspection Precision"], 1)])
    add_caption(doc, f"Table 8. Routing thresholds and test outcomes ({ALPHA} = 0.05).")
    make_table(doc, ["Model", "t_low*", "t_high*", "Val workload", "Val MCFR", "Test workload", "Test MCFR", "Urgent precision"], rows, font_size=6.8)
    add_paragraph(doc, "Random Forest gives the strongest operational routing result: 71.9% of test cases are auto-cleared while 0.35% of positives are routed to auto-clear. Several other models are conservatively routed to urgent inspection because the validation safety constraint cannot be satisfied with nonzero auto-clear workload.")
    add_paragraph(doc, "Figure 8 turns probability predictions into the three operational categories. The result is intentionally conservative for several models: if the validation constraint cannot be satisfied while auto-clearing cases, the threshold search produces little or no workload reduction. This is not a failure of the routing code. It is a useful diagnostic result because it shows that some calibrated models do not create a validation-safe low-risk region under the chosen alpha.")
    add_paragraph(doc, "Figure 9 shows the safety-workload trade-off directly. Random Forest achieves substantial workload reduction at alpha values above the strictest setting, while other models remain conservative. This sensitivity plot is more informative than reporting only alpha = 0.05 because it shows whether the routing result is stable or dependent on one arbitrary ceiling. Table 8 then reports the frozen alpha = 0.05 thresholds and test outcomes. The distinction between validation MCFR and test MCFR matters: the validation set is used for selection, while the test set is used only to estimate the final routing behavior.")

    add_paragraph(doc, "4.5. Explainability Diagnostics", "Heading 2")
    add_caption(doc, "Figure 10. Global feature importance for the structured XGBoost baseline.")
    doc.add_picture(str(FIG / "figure_10_shap_importance.png"), width=Inches(6.1))
    if (FIG / "figure_11_local_shap.png").exists():
        add_caption(doc, "Figure 11. Local SHAP diagnostics for true-positive, false-positive, false-negative, and true-negative structured XGBoost cases.")
        doc.add_picture(str(FIG / "figure_11_local_shap.png"), width=Inches(6.5))
    add_caption(doc, "Table 9. Top global SHAP feature importances.")
    rows = [[r["Feature"], fmt(r["Importance"], 4)] for _, r in imp.iterrows()]
    make_table(doc, ["Feature", "Importance"], rows, widths=[3.3, 1.2], font_size=8)
    add_paragraph(doc, "The top features are maintenance-history and sensor variables, especially time_since_last_maintenance, vibration_rms, and age. This is plausible under the simulator because maintenance timing, age, and vibration are direct or noisy consequences of degradation. SHAP is used here as a diagnostic audit, not as causal evidence.")
    add_paragraph(doc, "Figures 10 and 11, together with Table 9, are best interpreted as a leakage and plausibility audit. The strongest global feature is time_since_last_maintenance, followed by age and sensor variables. That pattern is plausible because the simulator makes maintenance timing and accumulated operation affect degradation. The local examples also show how the model can produce both true and false high-risk decisions from similar maintenance-history signals. This supports the use of explanations as diagnostic artifacts, but it does not prove causal validity. The XAI literature on predictive maintenance warns that feature-attribution methods are helpful for inspection but should not be treated as complete evidence that a model is correct [28-30].")

    add_paragraph(doc, "5. Discussion", "Heading 1")
    add_paragraph(doc, "The benchmark produces three main lessons. First, structured features dominate these short synthetic technician reports. Frozen MiniLM does not improve text-only performance over TF-IDF in a meaningful way, and fusion does not surpass Random Forest. This is a useful negative result because it discourages overclaiming from general-purpose embeddings in low-label maintenance text. Second, random splits can hide deployment risk: held-out-site FNR is much higher for tree and fusion models. Third, routing thresholds must be selected before test evaluation. The validation-constrained routing layer exposes whether a calibrated model is safe enough for auto-clear decisions.")
    add_paragraph(doc, "The first lesson should be interpreted in relation to the synthetic report generator. The text channel was designed to be sparse, short, and imperfect. Therefore, the weak text-only results do not contradict studies showing that real maintenance logs and work orders can be valuable [18-22]. Instead, they show that text value depends on report richness, coverage, and alignment with the prediction window. In this benchmark, the structured sensors and maintenance-history variables are more consistent signals of near-future maintenance requirement. A future benchmark variant could increase report richness, introduce longer work-order narratives, or add domain-adapted language models to test when text begins to add stronger value.")
    add_paragraph(doc, "The second lesson concerns evaluation protocol. The difference between random-split and held-out-site results is not a minor reporting detail. If only the random split were reported, the benchmark would suggest stronger generalization than the site-held-out experiment supports. This mirrors the broader concern that benchmark design can materially change conclusions [31,32]. For industrial O&M decision support, deployment often means applying a model to a new asset group, site, or operating condition. Held-out-site evaluation is therefore not an optional stress test; it is part of the trustworthiness question.")
    add_paragraph(doc, "The third lesson concerns thresholding. A default threshold of 0.5 is not an operational policy. Maintenance supervisors do not usually need a binary classifier alone; they need decisions about which cases can be routine, which cases require review, and which cases need urgent attention. The validation-selected routing layer makes that distinction explicit. It also prevents a common methodological error: choosing thresholds after seeing test performance. By selecting t_low and t_high on the validation set and then freezing them, the benchmark treats routing as part of the model-selection pipeline rather than as a post-hoc explanation of favorable test outcomes.")
    add_paragraph(doc, "The calibration and reject-option literature supports the general idea that probability quality and abstention matter in decision support [23-27], but this manuscript does not claim that alpha = 0.05 is a universal safety limit. The alpha sensitivity analysis is included precisely because the acceptable missed-fault ceiling is context-dependent. A real plant might set a stricter value for safety-critical equipment or a looser value for low-consequence routine monitoring. The benchmark provides the mechanism for studying that trade-off; it does not prescribe the operational value.")
    add_paragraph(doc, "The explanation results should also be read carefully. SHAP highlights features that the trained model used, not features that are physically causal in the real world. In this synthetic benchmark, maintenance timing, age, and vibration are plausible because the generator connects them to hidden degradation. In a real deployment, however, high importance for maintenance timing could reflect maintenance policy, data recording practices, or unmeasured site-specific confounding. The explanation layer is therefore useful for finding implausible reliance or leakage, but it is not a substitute for expert review and external validation.")
    add_paragraph(doc, "The result should not be read as evidence that Random Forest is universally superior. It is the best baseline in this synthetic configuration. The benchmark is meant to be reused by stronger methods, including semi-supervised, self-supervised, domain-adaptive, physics-informed, edge, or federated approaches.")
    add_paragraph(doc, "Overall, the benchmark is most useful as a controlled failure-mode laboratory. A researcher can change report ambiguity, sensor noise, missingness, label fraction, class imbalance, or site heterogeneity and then observe not only PR-AUC but also calibration, FNR, routing workload, and explanation stability. This is the main value of synthetic benchmarking: not to replace real industrial datasets, but to make specific stress conditions explicit and repeatable before moving to external validation.")

    add_paragraph(doc, "6. Limitations", "Heading 1")
    add_caption(doc, "Table 10. Limitations and mitigation strategies.")
    make_table(doc, ["Limitation", "Implication", "Mitigation or future work"], [
        ["Synthetic data only", "No real industrial deployment validation", "Validate on real or public condition-monitoring datasets before operational use"],
        ["Simplified degradation model", "Does not capture full asset physics", "Add physics-informed or asset-specific degradation models"],
        ["Template-generated reports", "Not equivalent to real technician language", "Use real work-order corpora or stronger language simulation"],
        ["Simulated routing", "No real supervisor behavior measured", "Run human-subject or expert-review studies"],
        ["General-purpose MiniLM", "Not a domain-specific maintenance language model", "Compare domain-adapted encoders in future work"],
    ], widths=[1.75, 2.25, 2.7], font_size=7.4)
    add_paragraph(doc, "Table 10 lists the main limitations, but several deserve emphasis. The first and most important limitation is that all records are synthetic. The benchmark can test internal consistency, robustness under controlled perturbation, and methodological discipline, but it cannot demonstrate real industrial deployment validity. Any operational claim would require external data, plant-specific expert review, and prospective or retrospective validation on real maintenance events.")
    add_paragraph(doc, "Second, the degradation process is simplified. It is designed to create useful statistical structure, not to reproduce material wear, thermodynamics, fluid dynamics, control logic, or component-specific failure physics. The generated sensors are noisy consequences of the hidden degradation process, but they are not produced by calibrated physical equations. This limitation is acceptable for a benchmark protocol paper only because it is stated clearly and because the code is released for inspection and modification.")
    add_paragraph(doc, "Third, technician reports are generated text, not real work-order language. Real maintenance text can include abbreviations, multilingual entries, domain-specific terminology, copied templates, missing context, and post-event updates. The generated reports capture only a limited version of this complexity. Therefore, the text-related findings should be used to compare benchmark modalities, not to make broad claims about all maintenance NLP.")
    add_paragraph(doc, "Fourth, the routing layer is simulated. It uses labels and calibrated probabilities to estimate auto-clear, review, and urgent-inspection categories, but it does not model human supervisor behavior, organizational risk tolerance, spare-parts constraints, cost of downtime, or the consequences of false urgent inspections. Future work should connect the routing protocol to cost models and expert evaluation.")

    add_paragraph(doc, "7. Conclusions", "Heading 1")
    add_paragraph(doc, "This article presents a reproducible synthetic low-data benchmark for trustworthy AI-based industrial O&M decision support. The benchmark evaluates predictive performance, calibration, robustness, held-out-site generalization, explanation diagnostics, and validation-constrained human-in-the-loop routing. The study shows that default thresholds can miss many positives, random splits can overstate site generalization, and routing constraints expose calibration weaknesses that aggregate metrics hide. The work is a benchmark and protocol contribution; it does not claim real industrial deployment validation.")
    add_paragraph(doc, "The main practical conclusion is that trustworthy O&M evaluation needs more than a model leaderboard. The benchmark shows how the same model can look acceptable under one metric and problematic under another. A model with strong ROC-AUC can still have high false negative rate at a default threshold; a model with reasonable chronological performance can degrade under held-out-site testing; a calibrated model can still fail to provide a safe auto-clear region under a missed-fault constraint. These observations are synthetic, but they reflect evaluation risks that the broader literature has repeatedly identified in low-data, shifted, and safety-relevant ML settings.")
    add_paragraph(doc, "Future work should extend the benchmark in three directions. The first is external validation using public or partner-provided industrial datasets. The second is richer simulation, including physics-informed degradation, multimodal time series, asset-specific failure modes, and more realistic work-order language. The third is decision support, including cost-aware routing, expert review, federated or edge deployment settings, and human-subject evaluation with maintenance supervisors. The current paper provides the reproducible baseline needed for those extensions.")

    add_paragraph(doc, "Back Matter", "Heading 1")
    add_paragraph(doc, "Supplementary Materials", "Heading 2")
    add_paragraph(doc, "The supplementary material consists of the public reproducibility repository, generated synthetic dataset, split manifests, result tables, high-resolution figures, graphical abstract, and the supplementary README included with this submission package.")
    add_paragraph(doc, "Funding", "Heading 2")
    add_paragraph(doc, "This research received no external funding.")
    add_paragraph(doc, "Author Contributions", "Heading 2")
    add_paragraph(doc, "Conceptualization, P.P.; methodology, P.P.; software, P.P.; validation, P.P.; formal analysis, P.P.; investigation, P.P.; data curation, P.P.; writing-original draft preparation, P.P.; writing-review and editing, P.P.; visualization, P.P.; project administration, P.P. The author has read and agreed to the submitted version of the manuscript.")
    add_paragraph(doc, "Data Availability Statement", "Heading 2")
    add_paragraph(doc, "The synthetic dataset generation code, benchmark configuration files, generated synthetic dataset, train/validation/test split files, label-scarcity split files, site-held-out split files, model-training scripts, threshold-selection scripts, evaluation scripts, result tables, and figure-generation artifacts are available at https://github.com/purohit0208/synthetic-low-data-om-benchmark. The dataset is fully synthetic and does not contain real industrial, personal, proprietary, or institution-owned operational data.")
    add_paragraph(doc, "Acknowledgments", "Heading 2")
    add_paragraph(doc, "The author thanks the developers of the open-source Python libraries used in the benchmark pipeline.")
    add_paragraph(doc, "Conflicts of Interest", "Heading 2")
    add_paragraph(doc, "The author declares no conflicts of interest.")
    add_paragraph(doc, "Generative AI Disclosure", "Heading 2")
    add_paragraph(doc, "During the preparation of this manuscript/study, the author used ChatGPT and Gemini for research planning, code drafting, manuscript structuring, and language refinement. The author reviewed and edited all outputs, verified the methods, results, and references, and takes full responsibility for the content of this publication.")
    add_paragraph(doc, "Synthetic-Data Statement", "Heading 2")
    add_paragraph(doc, "No real industrial data, personal data, confidential project data, proprietary institutional material, or third-party operational records were used. All records were synthetically generated for controlled methodological evaluation.")
    add_paragraph(doc, "References", "Heading 1")
    for ref in REFERENCES:
        add_paragraph(doc, ref)

    doc.save(SUB / "FINAL_MANUSCRIPT_MDPI_AI.docx")
    shutil.copy2(SUB / "FINAL_MANUSCRIPT_MDPI_AI.docx", ROOT / "manuscript" / "manuscript.docx")


REFERENCES = [
    "1. Hakami, A. Strategies for overcoming data scarcity, imbalance, and feature selection challenges in machine learning models for predictive maintenance. Scientific Reports 2024, 14, 9645. https://doi.org/10.1038/s41598-024-59958-9.",
    "2. Ramirez-Sanz, J.; Maestro-Prieto, J.; Arnaiz-Gonzalez, A.; Bustillo, A. Semi-supervised learning for industrial fault detection and diagnosis: A systemic review. ISA Transactions 2023, 143, 255-270. https://doi.org/10.1016/j.isatra.2023.09.027.",
    "3. Jalayer, M.; Kaboli, A.; Orsenigo, C.; Vercellis, C. Fault Detection and Diagnosis with Imbalanced and Noisy Data: A Hybrid Framework for Rotating Machinery. Machines 2022, 10, 237. https://doi.org/10.3390/machines10040237.",
    "4. Kafunah, J.; Ali, M.I.; Breslin, J.G. Handling Imbalanced Datasets for Robust Deep Neural Network-Based Fault Detection in Manufacturing Systems. Applied Sciences 2021, 11, 9783. https://doi.org/10.3390/app11219783.",
    "5. Swana, E.F.; Doorsamy, W.; Bokoro, P. Tomek Link and SMOTE Approaches for Machine Fault Classification with an Imbalanced Dataset. Sensors 2022, 22, 3246. https://doi.org/10.3390/s22093246.",
    "6. Shyalika, C.; Wickramarachchi, R.; El Kalach, F.; Harik, R.; Sheth, A. Evaluating the Role of Data Enrichment Approaches towards Rare Event Analysis in Manufacturing. Sensors 2024, 24, 5009. https://doi.org/10.3390/s24155009.",
    "7. Chen, C.; Fu, H.; Zheng, Y.; Tao, F.; Liu, Y. The advance of digital twin for predictive maintenance: The role and function of machine learning. Journal of Manufacturing Systems 2023, 71, 581-594. https://doi.org/10.1016/j.jmsy.2023.10.010.",
    "8. Hosamo, H.H.; Svennevig, P.R.; Svidt, K.; Han, D.; Nielsen, H.K. A Digital Twin predictive maintenance framework of air handling units based on automatic fault detection and diagnostics. Energy and Buildings 2022, 261, 111988. https://doi.org/10.1016/j.enbuild.2022.111988.",
    "9. Singh, R.R.; Bhatti, G.; Kalel, D.; Vairavasundaram, I.; Alsaif, F. Building a Digital Twin Powered Intelligent Predictive Maintenance System for Industrial AC Machines. Machines 2023, 11, 796. https://doi.org/10.3390/machines11080796.",
    "10. Yan, S.; Zhong, X.; Shao, H.; Ming, Y.; Liu, C.; Liu, B. Digital twin-assisted imbalanced fault diagnosis framework using subdomain adaptive mechanism and margin-aware regularization. Reliability Engineering & System Safety 2023, 239, 109522. https://doi.org/10.1016/j.ress.2023.109522.",
    "11. Mikolajewska, E.; Mikolajewski, D.; Mikolajczyk, T.; Paczkowski, T. Generative AI in AI-Based Digital Twins for Fault Diagnosis for Predictive Maintenance in Industry 4.0/5.0. Applied Sciences 2025, 15, 3166. https://doi.org/10.3390/app15063166.",
    "12. Moccardi, A.; Conte, C.; Ghosh, R.C.; Moscato, F. A Robust Conformal Framework for IoT-Based Predictive Maintenance. Future Internet 2025, 17, 244. https://doi.org/10.3390/fi17060244.",
    "13. Chen, L.; Li, Q.; Shen, C.; Zhu, J.; Wang, D.; Xia, M. Adversarial Domain-Invariant Generalization: A Generic Domain-Regressive Framework for Bearing Fault Diagnosis Under Unseen Conditions. IEEE Transactions on Industrial Informatics 2022, 18, 1790-1800. https://doi.org/10.1109/TII.2021.3078712.",
    "14. Asutkar, S.; Tallur, S. Deep transfer learning strategy for efficient domain generalisation in machine fault diagnosis. Scientific Reports 2023, 13, 6607. https://doi.org/10.1038/s41598-023-33887-5.",
    "15. Ragab, M.; Chen, Z.; Wu, M.; Foo, C.S.; Kwoh, C.K.; Yan, R.; Li, X. Contrastive Adversarial Domain Adaptation for Machine Remaining Useful Life Prediction. IEEE Transactions on Industrial Informatics 2021, 17, 5239-5249. https://doi.org/10.1109/TII.2020.3032690.",
    "16. Forest, F.; Fink, O. Calibrated Adaptive Teacher for Domain-Adaptive Intelligent Fault Diagnosis. Sensors 2024, 24, 7539. https://doi.org/10.3390/s24237539.",
    "17. Wen, B.; Xiao, M.; Wang, X.; Zhao, X.; Li, J.; Chen, X. Data-driven remaining useful life prediction based on domain adaptation. PeerJ Computer Science 2021, 7, e690. https://doi.org/10.7717/peerj-cs.690.",
    "18. Usuga-Cadavid, J.P.; Lamouri, S.; Grabot, B.; Fortin, A. Using deep learning to value free-form text data for predictive maintenance. International Journal of Production Research 2022, 60, 4548-4575. https://doi.org/10.1080/00207543.2021.1951868.",
    "19. Sundaram, S.; Zeid, A. Technical language processing for Prognostics and Health Management: applying text similarity and topic modeling to maintenance work orders. Journal of Intelligent Manufacturing 2025, 36, 1637-1657. https://doi.org/10.1007/s10845-024-02323-4.",
    "20. Giordano, V.; Fantoni, G. Decomposing maintenance actions into sub-tasks using natural language processing: A case study in an Italian automotive company. Computers in Industry 2024, 160, 104186. https://doi.org/10.1016/j.compind.2024.104186.",
    "21. Pavlopoulos, J.; Romell, A.; Curman, J.; Steinert, O.; Lindgren, T.; Borg, M.; Randl, K. Automotive fault nowcasting with machine learning and natural language processing. Machine Learning 2024, 113, 843-861. https://doi.org/10.1007/s10994-023-06398-7.",
    "22. Zhou, J.; Guo, Y.; Yang, Z.; Yang, J.; An, Z.; Li, K.; McLoone, S. T2MFDF: An LLM-Enhanced Multimodal Fault Diagnosis Framework Integrating Time-Series and Textual Data. IEEE Transactions on Instrumentation and Measurement 2025, 74, 1-11. https://doi.org/10.1109/TIM.2025.3583374.",
    "23. Silva Filho, T.; Song, H.; Perello-Nieto, M.; Santos-Rodriguez, R.; Kull, M.; Flach, P. Classifier calibration: a survey on how to assess and improve predicted class probabilities. Machine Learning 2023, 112, 3211-3260. https://doi.org/10.1007/s10994-023-06336-7.",
    "24. Hendrickx, K.; Perini, L.; Van der Plas, D.; Meert, W.; Davis, J. Machine learning with a reject option: a survey. Machine Learning 2024, 113, 3073-3110. https://doi.org/10.1007/s10994-024-06534-x.",
    "25. Hasan, M.M.; Abdar, M.; Khosravi, A.; Aickelin, U.; Lio, P.; Hossain, I.; Rahman, A.; Nahavandi, S. Survey on Leveraging Uncertainty Estimation Toward Trustworthy Deep Neural Networks: The Case of Reject Option and Post-Training Processing. ACM Computing Surveys 2025, 57, 236. https://doi.org/10.1145/3727633.",
    "26. Sayin, B.; Zoppi, T.; Marchini, N.; Khokhar, F.A.; Passerini, A. Bringing Machine Learning Classifiers Into Critical Cyber-Physical Systems: A Matter of Design. IEEE Access 2025, 13, 94858-94877. https://doi.org/10.1109/ACCESS.2025.3568501.",
    "27. Perez-Cerrolaza, J.; Abella, J.; Borg, M.; Donzella, C.; Cerquides, J.; Cazorla, F.J.; Englund, C.; Tauber, M.; Nikolakopoulos, G.; Flores, J.L. Artificial Intelligence for Safety-Critical Systems in Industrial and Transportation Domains: A Survey. ACM Computing Surveys 2024, 56, 176. https://doi.org/10.1145/3626314.",
    "28. Cummins, L.; Sommers, A.; Ramezani, S.B.; Mittal, S.; Jabour, J.; Seale, M.; Rahimi, S. Explainable Predictive Maintenance: A Survey of Current Methods, Challenges and Opportunities. IEEE Access 2024, 12, 57574-57602. https://doi.org/10.1109/ACCESS.2024.3391130.",
    "29. Brusa, E.; Cibrario, L.; Delprete, C.; Di Maggio, L.G. Explainable AI for Machine Fault Diagnosis: Understanding Features' Contribution in Machine Learning Models for Industrial Condition Monitoring. Applied Sciences 2023, 13, 2038. https://doi.org/10.3390/app13042038.",
    "30. Gawde, S.; Patil, S.; Kumar, S.; Kamat, P.; Kotecha, K.; Abraham, A. Explainable Predictive Maintenance of Rotating Machines Using LIME, SHAP, PDP, ICE. IEEE Access 2024, 12, 31743-31770. https://doi.org/10.1109/ACCESS.2024.3367110.",
    "31. Hendriks, J.; Dumond, P.; Knox, D.A. Towards better benchmarking using the CWRU bearing fault dataset. Mechanical Systems and Signal Processing 2022, 169, 108732. https://doi.org/10.1016/j.ymssp.2021.108732.",
    "32. Barnard, A.S. BenchMake: turn any scientific data set into a reproducible benchmark. Machine Learning: Science and Technology 2025, 6, 030502. https://doi.org/10.1088/2632-2153/adf810.",
    "33. McDermott, M.B.A.; Wang, S.; Marinsek, N.; Ranganath, R.; Foschini, L.; Ghassemi, M. Reproducibility in Machine Learning and Healthcare: How far do we have to go? Science Translational Medicine 2021, 13, eabb1655. https://doi.org/10.1126/scitranslmed.abb1655.",
]


def build_cover_letter():
    doc = Document()
    set_doc_styles(doc)
    add_paragraph(doc, "Cover Letter", "Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "22 May 2026")
    add_paragraph(doc, "Dear Editors,")
    add_paragraph(doc, "Please consider the manuscript titled \"A Synthetic Low-Data Benchmark for Trustworthy AI-Based Industrial Operation and Maintenance Decision Support\" for publication as an Article in AI, for the special issue \"AI for Industrial Operation and Maintenance: Recognition Challenges with Limited Data Condition.\"")
    add_paragraph(doc, "The manuscript contributes a reproducible synthetic benchmark and evaluation protocol for AI-based industrial O&M decision support under scarce labels, rare maintenance positives, noisy structured observations, imperfect technician reports, and cross-site domain shift. It evaluates structured-only, text-only, and fusion baselines; reports predictive performance, calibration, robustness, and explainability diagnostics; and introduces validation-set-selected human-in-the-loop routing thresholds under a missed-fault constraint.")
    add_paragraph(doc, "The work fits the special issue because it directly addresses AI-driven recognition challenges in industrial operation and maintenance under limited-data conditions, including sparse labels, rare fault events, synthetic data generation, noisy sensor observations, transfer/domain shift, uncertainty calibration, explainability, and human-supervisor routing.")
    add_paragraph(doc, "All data used in the study are synthetically generated. No real industrial data, personal data, confidential project data, proprietary institutional material, or third-party operational records were used.")
    add_paragraph(doc, "The reproducibility package is publicly available at https://github.com/purohit0208/synthetic-low-data-om-benchmark.")
    add_paragraph(doc, "Required statements: I confirm that neither the manuscript nor any parts of its content are currently under consideration for publication with, or published in, another journal. I confirm that the author has approved the manuscript and agrees with its submission to AI.")
    add_paragraph(doc, "Sincerely,")
    add_paragraph(doc, "Parth Purohit")
    add_paragraph(doc, "Independent Researcher")
    add_paragraph(doc, "Correspondence: To be supplied in the MDPI submission system.")
    doc.save(SUB / "COVER_LETTER_MDPI_AI.docx")


def write_submission_docs():
    SUPP.mkdir(parents=True, exist_ok=True)
    (SUB / "SUBMISSION_PACKAGE_CHECKLIST.md").write_text(
        "# MDPI AI Special Issue Submission Checklist\n\n"
        "Target journal: AI (MDPI)\n\n"
        "Special issue: AI for Industrial Operation and Maintenance: Recognition Challenges with Limited Data Condition\n\n"
        "Official special issue URL: https://www.mdpi.com/journal/ai/special_issues/60Q9P80Y81\n\n"
        "Core upload files prepared:\n\n"
        "- FINAL_MANUSCRIPT_MDPI_AI.docx\n"
        "- COVER_LETTER_MDPI_AI.docx\n"
        "- GRAPHICAL_ABSTRACT_MDPI_AI.png (optional but prepared; meets MDPI minimum size guidance)\n"
        "- figures_high_res.zip (optional high-resolution figure source package)\n"
        "- supplementary/SUPPLEMENTARY_MATERIALS_README.md\n\n"
        "Before online submission:\n\n"
        "- Confirm the author affiliation, correspondence email, ORCID, and biography in the MDPI submission system.\n"
        "- Select Article as the manuscript type.\n"
        "- Select the special issue named above during submission.\n"
        "- Confirm the cover-letter required statements are still true.\n"
        "- Confirm the expanded manuscript length remains appropriate for the target article type.\n"
        "- Confirm the 33 cited references are the intended local literature-review papers; the missing 31.pdf paper is not cited.\n"
        "- Upload only files you are allowed to distribute; do not upload the local literature_review folder or downloaded third-party PDFs.\n",
        encoding="utf-8",
    )
    (SUPP / "SUPPLEMENTARY_MATERIALS_README.md").write_text(
        "# Supplementary Materials README\n\n"
        "This supplementary package supports the manuscript titled \"A Synthetic Low-Data Benchmark for Trustworthy AI-Based Industrial Operation and Maintenance Decision Support\".\n\n"
        "Repository: https://github.com/purohit0208/synthetic-low-data-om-benchmark\n\n"
        "Contents expected in the public repository:\n\n"
        "- synthetic data generator and leakage audit;\n"
        "- configuration and metadata files;\n"
        "- generated synthetic dataset in CSV/Parquet;\n"
        "- train/validation/test, label-scarcity, random, and site-held-out split files;\n"
        "- model-training, calibration, routing, evaluation, and explanation scripts;\n"
        "- result tables and publication figures.\n\n"
        "Data boundary: all records are synthetic. No real industrial data, personal data, confidential project data, proprietary institutional material, or third-party operational records are included.\n",
        encoding="utf-8",
    )
    (SUB / "SPECIAL_ISSUE_ALIGNMENT.md").write_text(
        "# Special Issue Alignment\n\n"
        "The manuscript is aligned with the MDPI AI special issue \"AI for Industrial Operation and Maintenance: Recognition Challenges with Limited Data Condition\" because it addresses low-data industrial O&M recognition challenges through a controlled synthetic benchmark.\n\n"
        "Direct alignment points:\n\n"
        "- rare failure events and sparse labels;\n"
        "- synthetic data generation;\n"
        "- noisy sensor measurements;\n"
        "- transfer/domain shift across simulated plants;\n"
        "- uncertainty calibration and explainability;\n"
        "- human-in-the-loop routing for maintenance decision support.\n\n"
        "Scope boundary: the paper is an original benchmark and evaluation-protocol article using simulated industrial data. It does not claim real industrial deployment validation.\n",
        encoding="utf-8",
    )
    (SUB / "SUBMISSION_METADATA.md").write_text(
        "# Submission Metadata\n\n"
        "Use this file as a consistency aid when completing the MDPI submission form. Verify all personal and institutional details before submission.\n\n"
        "Journal: AI (MDPI)\n\n"
        "Special issue: AI for Industrial Operation and Maintenance: Recognition Challenges with Limited Data Condition\n\n"
        "Article type: Article\n\n"
        "Title: A Synthetic Low-Data Benchmark for Trustworthy AI-Based Industrial Operation and Maintenance Decision Support\n\n"
        "Author: Parth Purohit\n\n"
        "Affiliation: Independent Researcher\n\n"
        "Corresponding author email: To be supplied in the MDPI submission system.\n\n"
        "Keywords: predictive maintenance; industrial operation and maintenance; synthetic benchmark; low-data learning; domain shift; uncertainty calibration; human-in-the-loop decision support; technician reports; explainable AI; reproducibility\n\n"
        "Abstract:\n\n"
        "Industrial operation and maintenance (O&M) AI is difficult to evaluate when labeled failures are scarce, observations are noisy, technician reports are incomplete, and deployment sites differ from training sites. This article presents a reproducible synthetic benchmark for these conditions in a generic multi-site industrial setting. The generator simulates 480 assets across six plants over 365 days, producing 175,200 shift-level records with hidden degradation, noisy structured observations, imperfect technician reports, leakage-audited features, chronological splits, random splits, and site-held-out splits. Seven calibrated baselines are evaluated: structured-only classifiers, TF-IDF plus Logistic Regression, frozen MiniLM sentence embeddings, and structured-plus-text fusion. Random Forest obtains the highest chronological-test PR-AUC (0.398), but at the default 0.5 threshold still misses 58.5% of positives. Random-split evaluation overstates generalization: Random Forest PR-AUC falls from 0.708 under random splitting to 0.439 under held-out-site testing, and false negative rate increases from 0.471 to 0.815. Validation-set routing thresholds can reduce workload, but only for models whose low-risk calibration tails satisfy the missed-fault constraint. The study is fully synthetic and does not claim real industrial deployment validation.\n\n"
        "Funding statement: This research received no external funding.\n\n"
        "Data availability statement: The synthetic dataset generation code, benchmark configuration files, generated synthetic dataset, train/validation/test split files, label-scarcity split files, site-held-out split files, model-training scripts, threshold-selection scripts, evaluation scripts, result tables, and figure-generation artifacts are available at https://github.com/purohit0208/synthetic-low-data-om-benchmark. The dataset is fully synthetic and does not contain real industrial, personal, proprietary, or institution-owned operational data.\n\n"
        "Conflict of interest statement: The author declares no conflicts of interest.\n\n"
        "Generative AI disclosure: During the preparation of this manuscript/study, the author used ChatGPT and Gemini for research planning, code drafting, manuscript structuring, and language refinement. The author reviewed and edited all outputs, verified the methods, results, and references, and takes full responsibility for the content of this publication.\n\n"
        "Synthetic-data statement: No real industrial data, personal data, confidential project data, proprietary institutional material, or third-party operational records were used. All records were synthetically generated for controlled methodological evaluation.\n",
        encoding="utf-8",
    )
    zip_path = SUB / "figures_high_res.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(FIG.glob("*.png")):
            zf.write(p, p.name)


def main():
    generate_figures()
    build_docx()
    build_cover_letter()
    write_submission_docs()
    print(f"Built MDPI submission package in {SUB}")


if __name__ == "__main__":
    main()
