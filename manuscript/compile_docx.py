import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def add_paragraph_with_bold_runs(doc, text, style=None):
    """Adds a paragraph and parses **bold** and $math$ tags into runs."""
    p = doc.add_paragraph(style=style)
    # Split text by **bold** tags first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner_text = part[2:-2]
            # Check for math inside bold
            subparts = re.split(r'(\$.*?\$)', inner_text)
            for subpart in subparts:
                if subpart.startswith('$') and subpart.endswith('$'):
                    run = p.add_run(subpart[1:-1])
                    run.bold = True
                    run.italic = True
                    run.font.name = 'Consolas'
                else:
                    run = p.add_run(subpart)
                    run.bold = True
        else:
            # Check for math tags ($math$)
            subparts = re.split(r'(\$.*?\$)', part)
            for subpart in subparts:
                if subpart.startswith('$') and subpart.endswith('$'):
                    run = p.add_run(subpart[1:-1])
                    run.italic = True
                    run.font.name = 'Consolas'
                else:
                    p.add_run(subpart)
    return p

def docx_safe_image_path(image_path, script_dir):
    """Convert PNG figures to JPEG for more reliable headless DOCX/PDF rendering."""
    if not image_path.lower().endswith(".png"):
        return image_path

    output_dir = os.path.join(script_dir, "docx_images")
    os.makedirs(output_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(image_path))[0]
    output_path = os.path.join(output_dir, f"{base}.jpg")

    if os.path.exists(output_path) and os.path.getmtime(output_path) >= os.path.getmtime(image_path):
        return output_path

    with Image.open(image_path) as img:
        if img.mode in ("RGBA", "LA"):
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[-1])
            img = background
        else:
            img = img.convert("RGB")
        img.save(output_path, "JPEG", quality=94, optimize=True)

    return output_path

def main():
    md_path = "manuscript.md"
    docx_path = "manuscript.docx"
    
    # Resolve paths relative to script location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    repo_root = os.path.dirname(script_dir)
    
    if not os.path.isabs(md_path):
        md_path = os.path.join(script_dir, "manuscript.md")
        docx_path = os.path.join(script_dir, "manuscript.docx")
        
    print(f"Reading manuscript from {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Configure default styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    in_table = False
    table_lines = []
    
    for idx, line in enumerate(lines):
        striped = line.strip()
        
        # Detect Table lines
        if striped.startswith("|"):
            in_table = True
            table_lines.append(striped)
            continue
        elif in_table:
            # Table ended, process it
            in_table = False
            process_table(doc, table_lines)
            table_lines = []
            
        if not striped:
            continue
            
        # Detect block equations
        if striped.startswith("$$") and striped.endswith("$$"):
            eq_text = striped[2:-2].strip()
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_eq = p_eq.add_run(eq_text)
            run_eq.italic = True
            run_eq.font.name = 'Consolas'
            run_eq.font.size = Pt(11)
            # Add spacing
            p_eq.paragraph_format.space_before = Pt(6)
            p_eq.paragraph_format.space_after = Pt(6)
            continue
            
        # Detect images
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', striped)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve relative image path
            if not os.path.isabs(img_path):
                full_img_path = os.path.abspath(os.path.join(repo_root, img_path))
            else:
                full_img_path = img_path
                
            if os.path.exists(full_img_path):
                full_img_path = docx_safe_image_path(full_img_path, script_dir)
                print(f"Embedding image: {full_img_path}")
                try:
                    # Add picture (centered)
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    run_img = p_img.add_run()
                    run_img.add_picture(full_img_path, width=Inches(5.5))
                    
                    # Add caption (centered, small, italicized, grey)
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run(caption)
                    run_cap.italic = True
                    run_cap.font.size = Pt(9.5)
                    run_cap.font.color.rgb = RGBColor(100, 100, 100)
                except Exception as e:
                    print(f"Error embedding image {full_img_path}: {e}")
            else:
                print(f"Warning: Image file not found at: {full_img_path}")
            continue
            
        # Headers
        if striped.startswith("#"):
            level = len(striped) - len(striped.lstrip('#'))
            title_text = striped.lstrip('#').strip()
            
            # Stylize headings
            if level == 1:
                p = doc.add_heading(title_text, level=1)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(15)
                p.style.font.bold = True
                p.style.font.color.rgb = RGBColor(0, 91, 148)  # Deep blue primary
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            elif level == 2:
                p = doc.add_heading(title_text, level=2)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(13)
                p.style.font.bold = True
                p.style.font.color.rgb = RGBColor(51, 51, 51)  # Charcoal secondary
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
            elif level == 3:
                p = doc.add_heading(title_text, level=3)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11.5)
                p.style.font.bold = True
                p.style.font.color.rgb = RGBColor(85, 85, 85)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_heading(title_text, level=4)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
                p.style.font.bold = True
                p.style.font.italic = True
                p.style.font.color.rgb = RGBColor(100, 100, 100)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
            continue
            
        # Lists
        if striped.startswith("* ") or striped.startswith("- "):
            list_text = striped[2:].strip()
            add_paragraph_with_bold_runs(doc, list_text, style='List Bullet')
            continue
            
        # Numbered list items: Parse with literal prefix to avoid Word's automatic list linking bug
        num_list_match = re.match(r'^(\d+\.\s)(.*)', striped)
        if num_list_match:
            prefix = num_list_match.group(1)
            list_text = num_list_match.group(2)
            # Add as normal paragraph with manual left indent and literal prefix
            p = add_paragraph_with_bold_runs(doc, f"{prefix}{list_text}")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            continue
            
        # Horizontal rule
        if striped == "---":
            # Add a bottom border visual separator
            p = doc.add_paragraph()
            p_border = OxmlElement('w:pBdr')
            bottom_border = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>')
            p_border.append(bottom_border)
            p._p.get_or_add_pPr().append(p_border)
            continue
            
        # Standard paragraph
        add_paragraph_with_bold_runs(doc, striped)
        
    # In case file ended with a table
    if in_table and table_lines:
        process_table(doc, table_lines)
        
    print(f"Saving Word document to {docx_path}...")
    doc.save(docx_path)
    print("Compilation successful!")

def process_table(doc, lines):
    """Parses Markdown table lines and appends a Word table."""
    parsed_rows = []
    for line in lines:
        cells = [c.strip() for c in line.split("|")]
        if len(cells) > 1:
            if cells[0] == "":
                cells = cells[1:]
            if cells[-1] == "":
                cells = cells[:-1]
            parsed_rows.append(cells)
            
    if not parsed_rows:
        return
        
    # Check if second row is divider e.g. '---', ':---:', '---:'
    has_divider = False
    if len(parsed_rows) > 1:
        first_cell_divider = parsed_rows[1][0]
        if all(char in '-: ' for char in first_cell_divider) and len(first_cell_divider) > 0:
            has_divider = True
            
    header_row = parsed_rows[0]
    data_rows = parsed_rows[2:] if has_divider else parsed_rows[1:]
    
    num_cols = len(header_row)
    num_rows = len(data_rows) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Shading Accent 1'
    
    # Write and format headers
    hdr_cells = table.rows[0].cells
    for col_idx, col_name in enumerate(header_row):
        hdr_cells[col_idx].text = col_name
        set_cell_background(hdr_cells[col_idx], "005B94") # Theme primary deep blue
        # Bold and white font for headers
        for p in hdr_cells[col_idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
    # Write data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        # Shading for alternating rows
        bg_color = "F5F7FA" if row_idx % 2 == 0 else "FFFFFF"
        
        for col_idx in range(num_cols):
            val = row_data[col_idx] if col_idx < len(row_data) else ""
            row_cells[col_idx].text = val
            set_cell_background(row_cells[col_idx], bg_color)
            for p in row_cells[col_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9.5)
                
    # Space after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)

if __name__ == "__main__":
    main()
