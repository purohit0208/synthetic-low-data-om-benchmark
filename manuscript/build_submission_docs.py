from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
SUBMISSION_DIR = ROOT / "submission" / "mdpi_ai_2026_05_22"


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)


def build_cover_letter() -> Path:
    doc = Document()
    style_doc(doc)

    doc.add_heading("Cover Letter", level=1)
    add_paragraph(doc, "22 May 2026")
    add_paragraph(doc, "")
    add_paragraph(doc, "Dear Editors,")
    add_paragraph(
        doc,
        "Please consider the manuscript titled \"A Synthetic Low-Data Benchmark for Trustworthy AI-Based Industrial Operation and Maintenance Decision Support\" for publication as an Article in AI, for the special issue \"AI for Industrial Operation and Maintenance: Recognition Challenges with Limited Data Condition.\"",
    )
    add_paragraph(
        doc,
        "The manuscript contributes a reproducible synthetic benchmark and evaluation protocol for AI-based industrial O&M decision support under scarce labels, rare maintenance positives, noisy structured observations, imperfect technician reports, and cross-site domain shift. It evaluates structured-only, text-only, and fusion models; reports predictive performance, calibration, robustness, and explainability diagnostics; and introduces validation-set-selected human-in-the-loop routing thresholds under a missed-fault constraint.",
    )
    add_paragraph(
        doc,
        "The work fits the special issue because it directly addresses AI-driven industrial operation and maintenance under limited-data conditions, including sparse labels, rare fault events, synthetic data generation, noisy sensor observations, domain shift, uncertainty calibration, explainability, and human-supervisor decision routing.",
    )
    add_paragraph(
        doc,
        "All data used in the study are synthetically generated. No real industrial data, personal data, confidential project data, proprietary institutional material, or third-party operational records were used.",
    )
    add_paragraph(
        doc,
        "The reproducibility package is publicly available at https://github.com/purohit0208/synthetic-low-data-om-benchmark, release v1.0.0-mdpi-ai-submission.",
    )
    add_paragraph(
        doc,
        "Required statements: I confirm that neither the manuscript nor any parts of its content are currently under consideration for publication with, or published in, another journal. I confirm that the author has approved the manuscript and agrees with its submission to AI.",
        bold_prefix="Required statements:",
    )
    add_paragraph(doc, "Sincerely,")
    add_paragraph(doc, "Parth Purohit")
    add_paragraph(doc, "Independent Researcher")

    output = SUBMISSION_DIR / "COVER_LETTER_MDPI_AI.docx"
    doc.save(output)
    return output


def build_submission_checklist() -> Path:
    text = """# MDPI AI Submission Package Checklist

Prepared: 22 May 2026

Target journal: AI (MDPI)
Special issue: AI for Industrial Operation and Maintenance: Recognition Challenges with Limited Data Condition

Files prepared:

- FINAL_MANUSCRIPT_MDPI_AI.docx
- COVER_LETTER_MDPI_AI.docx
- SUBMISSION_PACKAGE_CHECKLIST.md

Public repository:

- https://github.com/purohit0208/synthetic-low-data-om-benchmark
- Release: v1.0.0-mdpi-ai-submission

Notes before online submission:

- Enter the author's real correspondence email in the MDPI submission system.
- Confirm that "Independent Researcher" is the correct affiliation. If a university or institutional affiliation should be used, update it before submission.
- Upload the cover letter when requested by the MDPI submission system.
- The generated synthetic dataset and reproducibility files are available through GitHub, so they do not need to be uploaded as a large supplementary file unless the editor specifically requests it.
"""
    output = SUBMISSION_DIR / "SUBMISSION_PACKAGE_CHECKLIST.md"
    output.write_text(text, encoding="utf-8")
    return output


def main() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    manuscript_src = MANUSCRIPT_DIR / "manuscript.docx"
    manuscript_dest = SUBMISSION_DIR / "FINAL_MANUSCRIPT_MDPI_AI.docx"
    copy2(manuscript_src, manuscript_dest)
    build_cover_letter()
    build_submission_checklist()
    print(f"Submission package written to {SUBMISSION_DIR}")


if __name__ == "__main__":
    main()
